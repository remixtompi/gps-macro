# -*- coding: utf-8 -*-
"""
Genera el informe explicativo de GPS_MACRO en formato Word (.docx).
Ejecutar: python generar_informe.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- Paleta ----------
AZUL = RGBColor(0x1F, 0x3A, 0x5F)
AZUL_CLARO = RGBColor(0x2E, 0x5E, 0x8C)
GRIS = RGBColor(0x44, 0x44, 0x44)
VERDE = RGBColor(0x15, 0x80, 0x3D)
NARANJA = RGBColor(0xB4, 0x53, 0x09)
ROJO = RGBColor(0xB0, 0x1B, 0x1B)

doc = Document()

# ---------- Estilos base ----------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = GRIS

def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def h1(text):
    p = doc.add_heading(level=1)
    r = p.add_run(text)
    r.font.color.rgb = AZUL
    r.font.size = Pt(18)
    r.bold = True
    return p

def h2(text):
    p = doc.add_heading(level=2)
    r = p.add_run(text)
    r.font.color.rgb = AZUL_CLARO
    r.font.size = Pt(14)
    r.bold = True
    return p

def h3(text):
    p = doc.add_heading(level=3)
    r = p.add_run(text)
    r.font.color.rgb = AZUL_CLARO
    r.font.size = Pt(12)
    r.bold = True
    return p

def para(text=None, bold=False, italic=False, size=11, color=GRIS, align=None, space_after=6):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        r.font.color.rgb = color
    return p

def bullet(text, bold_prefix=None, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.color.rgb = AZUL
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def numbered(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Number")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.color.rgb = AZUL
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def tabla(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htext)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(hdr[i], "1F3A5F")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
            run.font.color.rgb = GRIS
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t

def callout(title, text, color=AZUL_CLARO, fill="EAF1F8"):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    set_cell_bg(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = color
    r.font.size = Pt(11)
    p2 = cell.add_paragraph()
    r2 = p2.add_run(text)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = GRIS
    doc.add_paragraph()
    return t

# =====================================================================
# PORTADA
# =====================================================================
para()
para()
p = para("GPS MACRO", bold=True, size=40, color=AZUL, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("Sistema Cuantitativo de Navegacion Macro-Sectorial",
     bold=True, size=16, color=AZUL_CLARO, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("Top-Down: Macro -> Sector -> Precio",
     italic=True, size=13, color=GRIS, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

para("Informe explicativo y evaluacion tecnica",
     bold=True, size=13, color=GRIS, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("Que hace - Para que sirve - Como funciona - De donde saca los datos - "
     "Ventajas - Como aprovecharlo en el trading - Evaluacion honesta y mejoras",
     italic=True, size=11, color=GRIS, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)

para("Fecha del informe: 30 de mayo de 2026", size=11, color=GRIS, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("Documento de uso interno", italic=True, size=10, color=GRIS, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# =====================================================================
# 0. INDICE / RESUMEN EJECUTIVO
# =====================================================================
h1("1. Resumen ejecutivo")

para("GPS_MACRO es un sistema cuantitativo automatizado que responde a una sola "
     "pregunta de fondo para un inversor o trader: \"Dado el momento del ciclo economico, "
     "donde deberia estar puesto mi dinero?\". No es un robot que compra y vende; es una "
     "brujula (un GPS) que indica el sector del mercado con viento a favor y avisa cuando "
     "el entorno macroeconomico cambia.")

para("El sistema funciona en tres capas, de arriba hacia abajo (enfoque top-down):")
bullet("identifica en cual de las 4 fases del ciclo economico estamos (Expansion, "
       "Recalentamiento, Contraccion o Desaceleracion).", bold_prefix="Capa 1 - Macro: ")
bullet("puntua del 0 al 100 los 11 sectores de la bolsa de EE.UU. y senala los lideres "
       "y los rezagados.", bold_prefix="Capa 2 - Sector: ")
bullet("compara la salud real de la economia con el optimismo o pesimismo del mercado para "
       "detectar trampas (divergencias).", bold_prefix="Capa 3 - Divergencia: ")

para("Todo se publica solo, cada dia habil antes de la apertura del mercado de EE.UU., en "
     "un panel web (dashboard) gratuito alojado en GitHub Pages. El coste de operacion es "
     "cero: usa unicamente fuentes de datos publicas y gratuitas (la Reserva Federal y "
     "Yahoo Finance) e infraestructura gratuita de GitHub.")

callout("Idea clave del diseno",
        "El sistema te dice QUE mirar (que sector, que sub-sector). El CUANDO entrar y "
        "donde poner el stop siguen siendo decision tuya y de tu analisis tecnico. "
        "GPS_MACRO te pone en la carretera correcta; tu conduces.")

# =====================================================================
# 2. QUE ES Y PARA QUE SIRVE
# =====================================================================
h1("2. Que es GPS_MACRO y para que sirve")

h2("2.1. El problema que resuelve")
para("La mayoria de traders minoristas miran solo el grafico de un activo (analisis "
     "tecnico) y olvidan el contexto macroeconomico que mueve a TODO el mercado a la vez. "
     "El resultado tipico es comprar buenos graficos en sectores que la economia esta "
     "castigando, o ponerse defensivo justo cuando empieza una expansion. GPS_MACRO ataca "
     "ese punto ciego: ordena el tablero de juego ANTES de que elijas la ficha.")

h2("2.2. Para que sirve, en concreto")
bullet("Saber en que fase del ciclo economico estamos sin tener que leer decenas de "
       "informes economicos cada semana.")
bullet("Saber que sectores tienen el viento macroeconomico a favor y cuales en contra.")
bullet("Concentrar tu busqueda de oportunidades (setups tecnicos) en los 2-3 sectores "
       "lideres en lugar de en todo el mercado.")
bullet("Recibir avisos automaticos cuando cambia la fase del ciclo, cambia el sector "
       "lider, o cuando el mercado y la economia se desalinean (senal de cautela).")
bullet("Tener una vision objetiva y sin emociones: el mismo dato siempre produce el mismo "
       "resultado, no hay opiniones ni sesgos.")

h2("2.3. Para que NO sirve (limites honestos del alcance)")
bullet("No genera senales de compra/venta ni precios de entrada o stops.")
bullet("No opera intradia ni en marcos de tiempo cortos: su horizonte natural es de "
       "semanas a meses (rotacion de ciclo).")
bullet("No cubre acciones individuales, solo sectores y algunos sub-sectores via ETFs.")
bullet("No predice el futuro: describe el estado actual y su tendencia reciente.")

# =====================================================================
# 3. LOGICA DE FUNCIONAMIENTO
# =====================================================================
h1("3. Logica de funcionamiento (la filosofia top-down)")

para("El sistema imita como razona un gestor macro profesional, en cascada:")

tabla(
    ["Nivel", "Pregunta que responde", "Salida"],
    [
        ["1. Macro", "En que momento del ciclo economico estamos?", "1 de 4 fases + probabilidades"],
        ["2. Sector", "Que sectores se benefician de esa fase y ademas tienen fuerza tecnica?", "Ranking 0-100 de los 11 sectores"],
        ["3. Sub-sector", "Dentro del sector lider, que nicho es el mas fuerte?", "Ranking de sub-sectores (ETFs)"],
        ["4. Divergencia", "El mercado confirma o contradice a la economia?", "Senal de alineacion o trampa"],
    ],
    col_widths=[1.1, 3.2, 2.4],
)

para("La premisa academica detras es el modelo de rotacion sectorial de Sam Stovall "
     "(S&P): a lo largo del ciclo economico, el liderazgo del mercado rota de forma "
     "relativamente predecible entre sectores ciclicos, defensivos, sensibles a tipos y "
     "ligados a materias primas. GPS_MACRO codifica ese mapa y lo cruza con la fuerza "
     "tecnica real de cada sector, de modo que no se fia solo de la teoria: exige que el "
     "precio confirme.")

h2("3.1. Las 4 fases del ciclo")
para("El modelo coloca la economia en un plano de dos ejes: Crecimiento (eje horizontal) "
     "y Stress/Inflacion (eje vertical). Las 4 combinaciones definen las 4 fases:")

tabla(
    ["Fase", "Crecimiento", "Stress", "Caracter", "Sectores favorecidos", "Sectores a evitar"],
    [
        ["Expansion", "Alto", "Bajo", "Risk-on temprano", "Tecnologia, Consumo Disc., Comunicaciones, Industrial", "Defensivo, Utilities"],
        ["Recalentamiento", "Alto", "Alto", "Tarde en el ciclo", "Energia, Materiales, Industrial, Financiero", "Tecnologia, Utilities"],
        ["Contraccion", "Bajo", "Alto", "Modo defensivo", "Consumo Defensivo, Utilities, Salud", "Financiero, Consumo Disc., Materiales"],
        ["Desaceleracion", "Bajo", "Bajo", "Enfriamiento suave", "Salud, Consumo Def., Utilities, Inmobiliario", "Energia, Materiales, Industrial"],
    ],
    col_widths=[1.2, 0.9, 0.7, 1.1, 2.3, 1.8],
)

para("En lugar de declarar una fase de forma brusca (binaria), el sistema calcula la "
     "DISTANCIA de la situacion actual al centro de cada una de las 4 fases y reparte una "
     "probabilidad suave entre ellas (mediante una funcion softmax). Asi puedes ver, por "
     "ejemplo, \"70% Expansion, 20% Recalentamiento, 10% otras\", que es mucho mas honesto "
     "que un simple \"estamos en Expansion\".")

# =====================================================================
# 4. DE DONDE SACA LOS DATOS Y POR QUE ESOS
# =====================================================================
h1("4. De donde saca los datos y por que esos datos")

h2("4.1. Las dos fuentes (ambas gratuitas y publicas)")
bullet("la base de datos economica oficial de la Reserva Federal de St. Louis. "
       "Es la fuente de referencia mundial para datos macroeconomicos de EE.UU., "
       "fiable, gratuita y con una API estable. De aqui salen los indicadores economicos.",
       bold_prefix="FRED: ")
bullet("de aqui salen los precios y volumenes diarios de los ETFs sectoriales, "
       "el VIX y los futuros de cobre y oro. Si Yahoo falla, el sistema esta disenado para "
       "caer a Stooq como respaldo.", bold_prefix="Yahoo Finance: ")

h2("4.2. Los 9 indicadores que VOTAN la fase del ciclo y por que")
para("Cada indicador se eligio porque mide una dimension distinta y poco redundante del "
     "ciclo. La columna \"Eje\" indica si vota en Crecimiento o en Stress/Inflacion:")

tabla(
    ["Indicador", "Fuente", "Eje", "Por que se incluye"],
    [
        ["Empleo en Manufactura (proxy del PMI)", "FRED MANEMP", "Crecimiento", "Termometro adelantado de la actividad industrial."],
        ["Curva de tipos 10A-2A", "FRED T10Y2Y", "Crecimiento", "El predictor de recesiones mas fiable historicamente."],
        ["Peticiones de subsidio de paro (4 sem.)", "FRED IC4WSA", "Crecimiento", "Senal semanal y muy temprana del mercado laboral."],
        ["Spread High Yield (OAS)", "FRED BAMLH0A0HYM2", "Stress", "Mide el miedo en el credito corporativo (dinero real en riesgo)."],
        ["Indice de Condiciones Financieras (Chicago Fed)", "FRED NFCI", "Stress", "Resume si las condiciones financieras estan tensas o laxas."],
        ["Indice Coincidente / LEI proxy", "FRED USSLIND", "Crecimiento", "Estado actual agregado de la economia real."],
        ["Inflacion implicita 5 anos (breakeven)", "FRED T5YIE", "Stress", "Expectativas de inflacion del propio mercado de bonos."],
        ["Ratio Cobre/Oro", "Yahoo HG=F / GC=F", "Crecimiento", "Apetito por riesgo: cobre (industria) vs oro (refugio)."],
        ["VIX", "Yahoo ^VIX", "Stress", "El \"indice del miedo\" de la bolsa, en tiempo real."],
    ],
    col_widths=[2.5, 1.5, 1.0, 2.6],
)

callout("Por que ESTOS datos y no otros",
        "Mezclan deliberadamente tres velocidades: datos lentos pero fiables (empleo, LEI), "
        "datos de mercado en tiempo real (VIX, cobre/oro, spreads de credito) y datos "
        "adelantados (curva de tipos, peticiones de paro). Asi el modelo no depende de un "
        "solo tipo de senal y combina lo que la economia 'dice' con lo que el dinero "
        "'hace'. Ademas todos son gratuitos, diarios o semanales, y no se revisan tanto "
        "como el PIB, que llega tarde y se corrige meses despues.",
        color=NARANJA, fill="FBF0E5")

h2("4.3. Indicadores secundarios (informativos, no votan)")
para("Se muestran como contexto pero no afectan la clasificacion de fase: M2 (masa "
     "monetaria), Rendimiento real a 10 anos, Ventas minoristas, Permisos de construccion "
     "y el Indice del Dolar (DXY).")

h2("4.4. Como se procesa cada indicador (para que sea comparable)")
para("No se pueden sumar peras con manzanas (un VIX de 18 con un spread de 3,5%). Por eso "
     "cada serie pasa por la misma normalizacion robusta:")
numbered("se calcula su z-score recortado contra sus ultimos 10 anos usando mediana y MAD "
         "(desviacion absoluta mediana), que es resistente a valores extremos. Resultado "
         "acotado entre -3 y +3.", bold_prefix="Nivel: ")
numbered("se mide el cambio del indicador frente a hace 90 dias (tendencia reciente).",
         bold_prefix="Momentum: ")
numbered("score final = 70% nivel + 30% momentum, ajustado por el signo del indicador, "
         "acotado entre -1 y +1.", bold_prefix="Combinacion: ")
para("Despues se promedian los indicadores de cada eje para obtener un growth_score y un "
     "stress_score globales, que son las coordenadas que situan la economia en el plano de "
     "fases.")

# =====================================================================
# 5. EL SCORE SECTORIAL
# =====================================================================
h1("5. Como se puntuan los sectores (el motor de rotacion)")

para("Una vez conocida la fase, el sistema puntua del 0 al 100 cada uno de los 11 sectores "
     "GICS (representados por los ETFs SPDR: XLK, XLF, XLE, XLI, XLV, XLY, XLP, XLB, XLU, "
     "XLRE, XLC), comparandolos siempre contra el SPY (el mercado). El score combina 6 "
     "componentes, cada uno normalizado y con su peso:")

tabla(
    ["Componente", "Peso", "Que mide"],
    [
        ["Fuerza Relativa de Mansfield", "25%", "Si el sector lo hace mejor o peor que el mercado en 52 semanas."],
        ["Momentum 1M / 3M / 6M", "30%", "Aceleracion relativa reciente (pondera mas el corto: 50/30/20)."],
        ["Ranking transversal", "15%", "Posicion del sector frente a los otros 10."],
        ["Breadth (amplitud)", "10%", "Salud interna: precio vs medias de 50 y 200 dias."],
        ["Flujo de volumen (OBV)", "10%", "Si entra o sale dinero, via pendiente del OBV vs el mercado."],
        ["Alineacion con la fase macro", "10%", "Bonus si el sector esta favorecido por la fase; penalty si esta a evitar."],
    ],
    col_widths=[2.4, 0.8, 3.8],
)

para("El resultado se escala a 0-100 (50 = neutral). El sector con mayor score es el "
     "lider. Si ese lider supera un umbral de 70, el sistema baja un nivel mas y analiza "
     "sus sub-sectores (por ejemplo, si lidera Tecnologia, mira Semiconductores, Software "
     "y Ciberseguridad) para afinar donde esta exactamente la fuerza.")

callout("Detalle inteligente",
        "El score NO se fia solo de la teoria del ciclo. La alineacion con la fase macro "
        "pesa apenas un 10%; el 90% restante es fuerza tecnica real (precio, momentum, "
        "volumen). Esto evita comprar un sector 'que deberia ir bien segun el libro' pero "
        "que en la practica el mercado esta vendiendo.")

# =====================================================================
# 6. SENTIMIENTO Y DIVERGENCIAS
# =====================================================================
h1("6. Sentimiento de mercado y deteccion de divergencias")

para("La tercera capa calcula un indice de sentimiento del mercado (de -1 = panico a +1 = "
     "euforia) combinando cinco senales: tendencia del SPY (30%), VIX invertido (20%), "
     "spread High Yield invertido (25%), amplitud sectorial (15%) y apetito de riesgo "
     "medido como Consumo Discrecional frente a Consumo Defensivo (10%).")

para("Despues compara ese sentimiento del mercado con la salud macro real (growth menos "
     "stress). Cuando ambos se separan mas de un umbral, declara una divergencia:")
bullet("la economia esta debil pero el mercado sube eufórico. Senal de "
       "cautela: posible euforia insostenible.", bold_prefix="Mercado mejor que macro: ")
bullet("la economia es solida pero el mercado esta deprimido. "
       "Posible oportunidad / pesimismo exagerado.", bold_prefix="Macro mejor que mercado: ")
para("Las divergencias son justamente las situaciones donde mas traders caen en trampas, "
     "por eso el sistema las marca con un banner de aviso en el dashboard.")

# =====================================================================
# 7. ARQUITECTURA Y AUTOMATIZACION
# =====================================================================
h1("7. Arquitectura tecnica y automatizacion")

h2("7.1. El pipeline diario, paso a paso")
numbered("Descarga datos macro (FRED) y de mercado (Yahoo), con cache local de 6 horas y "
         "reintentos automaticos.")
numbered("Calcula la fase del ciclo y sus probabilidades.")
numbered("Puntua los 11 sectores y, si procede, los sub-sectores del lider.")
numbered("Calcula sentimiento y detecta divergencias.")
numbered("Actualiza el historico (un archivo JSON versionado en Git) y detecta cambios de "
         "fase, cambios de lider y divergencias para generar alertas.")
numbered("Renderiza un dashboard web estatico (HTML + graficos) con plantillas Jinja2.")
numbered("Hace commit y push automatico; GitHub Pages publica la nueva version.")

h2("7.2. La infraestructura (todo gratis)")
tabla(
    ["Pieza", "Tecnologia", "Coste"],
    [
        ["Motor de calculo", "Python (pandas, numpy)", "Gratis / open source"],
        ["Datos economicos", "FRED API", "Gratis (requiere API key)"],
        ["Datos de mercado", "Yahoo Finance (+ Stooq de respaldo)", "Gratis"],
        ["Automatizacion / cron", "GitHub Actions", "Gratis (repos publicos)"],
        ["Publicacion web", "GitHub Pages", "Gratis"],
        ["Almacenamiento historico", "JSON versionado en Git", "Gratis"],
    ],
    col_widths=[2.0, 2.8, 2.2],
)
para("El robot (GitHub Actions) se ejecuta de lunes a viernes a las 12:30 UTC "
     "(pre-apertura de EE.UU.) y los sabados a mediodia para un repaso semanal. El "
     "mantenimiento humano es practicamente nulo.")

h2("7.3. Decisiones de diseno destacables")
bullet("mismo dato de entrada produce siempre el mismo resultado. No hay "
       "IA generativa en el pipeline de decision; todo es auditable y reproducible.",
       bold_prefix="Determinista: ")
bullet("el historico vive como JSON en Git, por lo que esta "
       "versionado, es reversible y no genera costes de hosting.", bold_prefix="Sin base de datos: ")
bullet("si una fuente falla, el sistema cae a la alternativa o "
       "continua con los datos disponibles sin romperse.", bold_prefix="Tolerante a fallos: ")
bullet("incluye un modo de datos sinteticos para poder probar todo el "
       "pipeline sin internet ni claves.", bold_prefix="Testeable: ")

# =====================================================================
# 8. COMO LEER EL DASHBOARD / SACARLE PROVECHO
# =====================================================================
h1("8. Como aprovecharlo en el trading (flujo de trabajo)")

para("GPS_MACRO no sustituye tu metodo; lo potencia. Un flujo de trabajo recomendado:")

numbered("Mira la fase del ciclo y su probabilidad. Esto define tu sesgo general "
         "(ofensivo o defensivo).", bold_prefix="Contexto (1 min): ")
numbered("Lee el ranking sectorial. Quedate con el top 2-3. Ahi vive el dinero "
         "inteligente ahora mismo.", bold_prefix="Donde mirar: ")
numbered("Si el lider tiene score >= 70, baja al sub-sector mas fuerte para afinar la "
         "puntería.", bold_prefix="Afinar: ")
numbered("Revisa el banner de divergencia. Si hay aviso, reduce tamano o exige mas "
         "confirmacion.", bold_prefix="Filtro de riesgo: ")
numbered("Aplica TU analisis tecnico geometrico solo sobre los activos de esos sectores "
         "lideres para decidir el cuando entrar y donde poner el stop.",
         bold_prefix="Ejecucion (lo tuyo): ")

callout("Regla de oro",
        "GPS_MACRO responde DONDE y filtra el CUANDO-NO (cuando hay divergencia, cuidado). "
        "Tu sigues siendo el dueno del gatillo: el cuando-si y el stop son tu edge.")

h2("8.1. Como leer cada bloque del panel")
bullet("rojo/amarillo/azul; salta ante cambio de fase, divergencia o cambio "
       "de lider.", bold_prefix="Banner superior: ")
bullet("la fase actual y las probabilidades de las 4 fases.", bold_prefix="Gauge de fase: ")
bullet("valor, z-score (vs 10 anos), momentum 90d, score y mini-grafico por "
       "indicador.", bold_prefix="Tabla macro: ")
bullet("los 11 sectores con su score 0-100, ordenados visualmente.",
       bold_prefix="Heatmap sectorial: ")
bullet("top-3 resaltado y desglose de los componentes del score.",
       bold_prefix="Ranking detallado: ")
bullet("graficos de evolucion de fase, crecimiento, stress y sentimiento.",
       bold_prefix="Historico: ")

# =====================================================================
# 9. VENTAJAS Y BENEFICIOS
# =====================================================================
h1("9. Ventajas y beneficios")

h2("9.1. Ventajas tecnicas")
bullet("Coste cero, sin dependencias de pago.")
bullet("Totalmente automatizado y desatendido (corre solo cada dia).")
bullet("Determinista y auditable: sin cajas negras ni opiniones.")
bullet("Objetivo: elimina el sesgo emocional y la narrativa de las noticias.")
bullet("Multi-senal: combina economia, credito, volatilidad y precio, no un solo dato.")
bullet("Normalizacion robusta (mediana/MAD) resistente a outliers y a datos sucios.")
bullet("Tolerante a fallos de fuentes con cache, reintentos y respaldo.")
bullet("Calibrable: todos los pesos y umbrales viven en un solo archivo de configuracion.")

h2("9.2. Beneficios practicos para el trader")
bullet("Ahorra horas de lectura de informes macro cada semana.")
bullet("Concentra el esfuerzo en pocos sectores, mejorando la calidad de las decisiones.")
bullet("Evita pelear contra el ciclo (no comprar ciclicos en plena contraccion).")
bullet("Avisa de trampas (divergencias) antes de que te atrapen.")
bullet("Aporta disciplina y un marco repetible, no improvisacion.")

# =====================================================================
# 10. EVALUACION HONESTA
# =====================================================================
h1("10. Evaluacion honesta: que tan bueno es y que le falta")

para("Veredicto general: es un sistema solido, bien disenado y con una arquitectura "
     "profesional para su coste (cero). La logica top-down es correcta y academicamente "
     "fundada, la ingenieria es limpia y tolerante a fallos, y el alcance esta bien "
     "delimitado (no promete mas de lo que puede dar). Dicho esto, tiene limitaciones "
     "reales que conviene conocer.", bold=True, color=AZUL)

h2("10.1. Fortalezas")
bullet("Diseno conceptual correcto y honesto sobre su alcance.")
bullet("Codigo modular, limpio y mantenible.")
bullet("Resiliencia (cache, reintentos, fallback, modo sintetico).")
bullet("Cero coste y cero mantenimiento.")

h2("10.2. Limitaciones y puntos debiles (lo que le falta)")
bullet("los pesos (25/30/15/10/10/10), umbrales (fase >55%, sub-sector >=70) "
       "y los centros de las fases estan puestos por criterio experto, no validados "
       "estadisticamente. No sabemos si funcionan hasta probarlos.",
       bold_prefix="Sin backtest historico: ")
bullet("no hay ninguna medida de si las senales del sistema "
       "habrian generado rentabilidad o no. Es la carencia mas importante.",
       bold_prefix="Sin validacion de rendimiento: ")
bullet("breadth y volumen usan el precio del propio ETF como aproximacion, "
       "no los datos internos reales de las acciones que componen cada sector. Es una "
       "simplificacion razonable pero imperfecta.", bold_prefix="Proxies en lugar de datos reales: ")
bullet("solo cubre EE.UU. No mira Europa, emergentes, bonos, "
       "materias primas como clase de activo, ni cripto.", bold_prefix="Universo limitado: ")
bullet("no hay un test que compruebe que, ante una "
       "fase historica conocida (p.ej. 2008 o 2020), el sistema la clasifica bien.",
       bold_prefix="Sin validacion de la clasificacion de fase: ")
bullet("muchos datos FRED se revisan despues de "
       "publicarse; el sistema usa la ultima version, no la que se conocia ese dia, lo que "
       "puede inflar la calidad aparente si algun dia se hace backtest.",
       bold_prefix="No corrige sesgo de revision de datos: ")
bullet("si una serie FRED clave falla, el eje se promedia con menos "
       "indicadores sin avisar de la perdida de fiabilidad.", bold_prefix="Sin control de calidad de datos: ")
bullet("ningun aviso te llega salvo que entres a mirar la web.",
       bold_prefix="Sin alertas activas (push): ")

# =====================================================================
# 11. MEJORAS PROPUESTAS
# =====================================================================
h1("11. Mejoras propuestas (gratis y automatizables, tu y yo)")

para("Todas las siguientes mejoras se pueden implementar con herramientas gratuitas y de "
     "forma automatizada, sin coste de suscripcion. Las ordeno por relacion impacto/esfuerzo.",
     italic=True)

h2("11.1. Prioridad ALTA (alto impacto)")

h3("A. Modulo de backtest y validacion historica")
bullet("Descargar 15-20 anos de las mismas series (FRED y Yahoo dan historico gratis) y "
       "reconstruir la fase y el ranking sectorial dia a dia hacia atras.")
bullet("Medir si \"sobreponderar el top-3 sectorial\" habria batido al SPY, con metricas "
       "estandar (CAGR, max drawdown, Sharpe, hit-rate de las fases).")
bullet("Herramientas gratis: pandas + vectorbt o backtesting.py (open source). Lo podemos "
       "construir juntos sobre el codigo que ya existe.")
callout("Por que es la #1",
        "Sin esto, no sabemos si el sistema realmente aporta valor o solo 'se ve bien'. "
        "Es lo que convierte una herramienta interesante en una herramienta en la que "
        "puedes confiar tu dinero.", color=ROJO, fill="FBE9E9")

h3("B. Alertas push automaticas (Telegram / email)")
bullet("Cuando el pipeline detecte cambio de fase, cambio de lider o divergencia, enviar "
       "un mensaje automatico.")
bullet("Herramientas gratis: bot de Telegram (API gratuita) o email via GitHub Actions. "
       "Son ~20 lineas de codigo y un secret mas. Te enteras sin entrar a la web.")

h3("C. Validacion de calidad de datos")
bullet("Antes de calcular, comprobar que cada serie tiene datos recientes y suficientes; "
       "marcar el dashboard como \"fiabilidad reducida\" si falta algo importante.")

h2("11.2. Prioridad MEDIA")

h3("D. Backtest del sentimiento y de las divergencias")
bullet("Medir historicamente si las divergencias marcadas anticiparon caidas o rebotes. "
       "Permite calibrar el umbral (hoy fijado en 0,35 por criterio).")

h3("E. Ampliar el universo con datos gratuitos")
bullet("Anadir bonos (TLT), oro (GLD), dolar (UUP) y quizas Europa/emergentes (VGK, EEM) "
       "como contexto de clase de activo. Todo via Yahoo, gratis.")

h3("F. Datos internos reales de breadth")
bullet("Sustituir el proxy de amplitud por datos reales (p.ej. % de componentes sobre su "
       "media de 200) usando fuentes gratuitas o calculando sobre las holdings de cada ETF.")

h3("G. Estado de animo de noticias macro (NLP gratis)")
bullet("Leer titulares macro de fuentes RSS gratuitas y puntuar su tono con un modelo de "
       "sentimiento open source, como capa de contexto adicional (no de decision).")

h2("11.3. Prioridad BAJA (mejoras de comodidad)")
bullet("Exportacion automatica del estado diario a Google Sheets (API gratuita) para tu "
       "propio registro y analisis.")
bullet("Version movil / PWA del dashboard para consultarlo comodo desde el telefono.")
bullet("Comparador historico: \"como estaba el tablero hace 1, 3 y 6 meses\".")

h2("11.4. Resumen de la hoja de ruta sugerida")
tabla(
    ["#", "Mejora", "Impacto", "Esfuerzo", "Coste"],
    [
        ["1", "Backtest / validacion historica", "Muy alto", "Medio-alto", "0 EUR"],
        ["2", "Alertas push (Telegram/email)", "Alto", "Bajo", "0 EUR"],
        ["3", "Validacion de calidad de datos", "Alto", "Bajo", "0 EUR"],
        ["4", "Backtest de divergencias", "Medio", "Medio", "0 EUR"],
        ["5", "Ampliar universo (bonos/oro/intl)", "Medio", "Bajo", "0 EUR"],
        ["6", "Breadth con datos reales", "Medio", "Medio", "0 EUR"],
        ["7", "Sentimiento de noticias (NLP)", "Medio-bajo", "Medio", "0 EUR"],
        ["8", "Export a Google Sheets / PWA", "Bajo", "Bajo", "0 EUR"],
    ],
    col_widths=[0.4, 3.0, 1.1, 1.1, 0.9],
)

# =====================================================================
# 12. CONCLUSION
# =====================================================================
h1("12. Conclusion")

para("GPS_MACRO es un sistema bien concebido y bien construido que cumple exactamente lo "
     "que promete: ser una brujula macro-sectorial automatica, gratuita y objetiva que te "
     "dice donde poner el foco y te avisa cuando el entorno cambia. Su mayor virtud es la "
     "honestidad de su diseno: no pretende sustituir tu criterio tecnico, solo darle "
     "contexto.")

para("Su mayor carencia, hoy, es la falta de validacion historica (backtest): tenemos un "
     "buen mapa, pero todavia no hemos comprobado contra el pasado que lleve al tesoro. "
     "Esa es, con diferencia, la mejora prioritaria, y es perfectamente realizable entre "
     "los dos, con herramientas gratuitas, sobre el codigo que ya existe.")

para("Con el backtest y las alertas push anadidas, GPS_MACRO pasaria de ser una "
     "herramienta interesante a ser una pieza central y fiable de tu proceso de inversion "
     "top-down.", bold=True, color=AZUL)

para()
para("--- Fin del informe ---", italic=True, size=10, color=GRIS, align=WD_ALIGN_PARAGRAPH.CENTER)

# ---------- Guardar ----------
out = "GPS_MACRO_Informe_Explicativo.docx"
doc.save(out)
print("Documento generado:", out)
